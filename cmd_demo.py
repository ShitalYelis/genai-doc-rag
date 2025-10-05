# cmd_demo_improved.py
import os
import re
import sys
import pdfplumber
from docx import Document as DocxDocument
import tkinter as tk
from tkinter import filedialog
from typing import List, Dict

# LangChain imports (robust import fallback)
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document as LCDocument
    from langchain_community.vectorstores import Chroma
    from langchain_community.embeddings import SentenceTransformerEmbeddings
except Exception:
    # Fallback if older/newer package layout
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document as LCDocument
    from langchain.vectorstores import Chroma
    from langchain.embeddings import SentenceTransformerEmbeddings

from transformers import pipeline

# --------------------------
# Configuration
# --------------------------
MODEL_NAME = os.environ.get("RAG_LLM", "google/flan-t5-large")  # change to flan-t5-base if memory is limited
RETRIEVE_K = 10   # number of chunks to retrieve for RAG fallback
CONTEXT_CHAR_LIMIT = 4000  # limit context characters sent to LLM
CHUNK_SIZE = 800
CHUNK_OVERLAP = 100

# --------------------------
# Helpers: text extraction
# --------------------------
def extract_text(file_path: str) -> str:
    """Extract text from pdf/docx/txt. Returns empty string on failure."""
    file_path = str(file_path)
    if file_path.lower().endswith(".pdf"):
        try:
            with pdfplumber.open(file_path) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            return "\n".join(pages)
        except Exception as e:
            print(f"⚠️ PDF extraction failed for {file_path}: {e}")
            return ""
    elif file_path.lower().endswith(".docx"):
        try:
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            print(f"⚠️ DOCX extraction failed for {file_path}: {e}")
            return ""
    elif file_path.lower().endswith(".txt"):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            print(f"⚠️ TXT read failed for {file_path}: {e}")
            return ""
    else:
        return ""

# --------------------------
# Helpers: sentence splitting & simple search
# --------------------------
_SENTENCE_SPLIT_RE = re.compile(r'(?<=[\.\?\!\n])\s+')

def split_sentences(text: str) -> List[str]:
    text = text.replace("\r", " ")
    # keep sentences not too short
    sents = [s.strip() for s in _SENTENCE_SPLIT_RE.split(text) if s.strip()]
    return sents

def find_sentences_with_all_keywords(texts: Dict[str, str], keywords: List[str], max_results=10) -> List[Dict]:
    """Search all uploaded texts (per file) for sentences that contain all keywords (case-insensitive).
       Returns list of {source, sentence} (up to max_results)."""
    kws = [k.lower() for k in keywords if k.strip()]
    results = []
    for fname, full in texts.items():
        for sent in split_sentences(full):
            low = sent.lower()
            if all(k in low for k in kws):
                results.append({"source": fname, "sentence": sent})
                if len(results) >= max_results:
                    return results
    return results

def count_term_across_texts(texts: Dict[str, str], term: str) -> Dict:
    """Count occurrences of exact term (word-boundary, case-insensitive) across all texts.
       Returns {'total': int, 'by_file': {fname:count}}"""
    term_re = re.compile(r"\b" + re.escape(term) + r"\b", flags=re.IGNORECASE)
    by_file = {}
    total = 0
    for fname, txt in texts.items():
        c = len(term_re.findall(txt))
        by_file[fname] = c
        total += c
    return {"total": total, "by_file": by_file}

# --------------------------
# Start: file picker
# --------------------------
print("📤 Please select one or more documents (PDF / DOCX / TXT)...")
root = tk.Tk()
root.withdraw()
file_paths = filedialog.askopenfilenames(
    title="Select documents",
    filetypes=[("Documents", "*.pdf *.docx *.txt")]
)

if not file_paths:
    print("❌ No files selected. Exiting.")
    sys.exit(0)

# load texts per file
uploaded_texts: Dict[str, str] = {}
for p in file_paths:
    fname = os.path.basename(p)
    if fname.startswith("~$"):
        print(f"⚠️ Skipping MS Word temp file: {fname}")
        continue
    txt = extract_text(p)
    if txt and txt.strip():
        uploaded_texts[fname] = txt
        print(f"✅ Loaded: {fname} (chars: {len(txt)})")
    else:
        print(f"⚠️ Skipped (no readable text): {fname}")

if not uploaded_texts:
    print("❌ No valid documents to process. Exiting.")
    sys.exit(0)

# --------------------------
# Create LangChain Documents (with metadata)
# --------------------------
print("\n📑 Splitting documents into chunks...")
splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
lc_documents: List[LCDocument] = []
for fname, txt in uploaded_texts.items():
    chunks = splitter.split_text(txt)
    for i, c in enumerate(chunks):
        metadata = {"source": fname, "chunk_id": i}
        lc_documents.append(LCDocument(page_content=c, metadata=metadata))
print(f"✅ Created {len(lc_documents)} chunks from {len(uploaded_texts)} files")

# --------------------------
# Build embeddings + Chroma vector DB
# --------------------------
print("\n🔍 Creating embeddings and building vector store (Chroma) ...")
embedding_model = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
# store to memory (collection name)
vectordb = Chroma.from_documents(lc_documents, embedding_model, collection_name="genai_docs")
retriever = vectordb.as_retriever(search_kwargs={"k": RETRIEVE_K})
print("✅ Vector DB ready")

# --------------------------
# Load LLM pipeline (HuggingFace)
# --------------------------
print(f"\n🤖 Loading LLM pipeline: {MODEL_NAME} (this may take a while the first time)...")
try:
    hf_pipe = pipeline("text2text-generation", model=MODEL_NAME, max_new_tokens=512, temperature=0.1)
except Exception as e:
    print("❗ Failed to load model. If you have low RAM/CPU, change MODEL_NAME to 'google/flan-t5-base'.")
    print("Error:", e)
    sys.exit(1)
print("✅ LLM pipeline loaded")

# --------------------------
# Utility: Build context string from retrieved chunks (limited)
# --------------------------
def build_context_from_docs(docs: List[LCDocument], char_limit=CONTEXT_CHAR_LIMIT) -> str:
    ctx_parts = []
    current_len = 0
    for d in docs:
        meta = d.metadata or {}
        source = meta.get("source", "unknown")
        chunk_id = meta.get("chunk_id", "")
        piece = f"[{source} | chunk:{chunk_id}]\n{d.page_content.strip()}\n\n"
        piece_len = len(piece)
        # avoid exceeding char limit, but include at least one chunk
        if current_len + piece_len > char_limit and ctx_parts:
            break
        ctx_parts.append(piece)
        current_len += piece_len
    return "\n".join(ctx_parts).strip()

# --------------------------
# Main interactive loop: improved logic
# --------------------------
print("\n🎯 System ready. Ask questions about your uploaded documents. Type 'exit' to quit.\n")

while True:
    query = input("🔎 Your question: ").strip()
    if not query:
        continue
    if query.lower() in ("exit", "quit"):
        print("👋 Bye.")
        break

    # DETERMINE: is it a 'count' question?
    if re.search(r"\b(count|how many|total count|number of|total number)\b", query, flags=re.I):
        # Try detect target term inside quotes first
        m = re.search(r'["\'](.+?)["\']', query)
        term = None
        if m:
            term = m.group(1).strip()
        else:
            # Try extract Title Case phrases e.g. Evidence Description or multi-word capitalized
            caps = re.findall(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', query)
            if caps:
                term = caps[0].strip()
            else:
                # last resort: take words after 'for' or 'of'
                m2 = re.search(r'(?:for|of)\s+(.+)$', query, flags=re.I)
                if m2:
                    term = m2.group(1).strip().strip(' ?.')
        if not term:
            print("⚠️ Couldn't detect term to count from your question. Falling back to retrieval-based answer.")
        else:
            counts = count_term_across_texts(uploaded_texts, term)
            print(f"\n📊 Total occurrences of '{term}': {counts['total']}")
            for fname, c in counts["by_file"].items():
                if c:
                    print(f" - {fname}: {c}")
            # Also show example sentences if exists
            example_sents = find_sentences_with_all_keywords(uploaded_texts, [term], max_results=5)
            if example_sents:
                print("\n📄 Example contexts (where the term appears):")
                for ex in example_sents:
                    print(f"[{ex['source']}] {ex['sentence']}")
            print("-" * 80)
            continue

    # DETERMINE: is it a 'details' / extraction about a named form/section?
    if re.search(r'\b(detail|provide details|describe|provide information|give details)\b', query, flags=re.I):
        # Attempt to extract a phrase (similar to count logic)
        m = re.search(r'["\'](.+?)["\']', query)
        phrase = None
        if m:
            phrase = m.group(1).strip()
        else:
            caps = re.findall(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', query)
            if caps:
                phrase = caps[0].strip()
            else:
                m2 = re.search(r'(?:about|for|regarding)\s+(.+)$', query, flags=re.I)
                if m2:
                    phrase = m2.group(1).strip().strip(' ?.')
        if phrase:
            hits = find_sentences_with_all_keywords(uploaded_texts, phrase.split(), max_results=10)
            if hits:
                print("\n📂 Found the following sentences that match your request:\n")
                for h in hits:
                    print(f"[{h['source']}] {h['sentence']}\n")
                print("-" * 80)
                continue
            else:
                print("⚠️ Could not find direct sentence matches for that phrase. Falling back to RAG retrieval.")
        else:
            print("⚠️ Could not identify the phrase you want details for. Falling back to RAG retrieval.")

    # GENERAL: use vector retrieval + strict prompt
    retrieved_docs = vectordb.similarity_search(query, k=RETRIEVE_K)
    if not retrieved_docs:
        print("⚠️ No relevant content found by semantic search.")
        print("-" * 80)
        continue

    # Show top 3 snippets to the user for transparency
    print("\n📂 Top relevant document snippets (for transparency):\n")
    for i, d in enumerate(retrieved_docs[:3], 1):
        src = d.metadata.get("source", "unknown")
        cid = d.metadata.get("chunk_id", "")
        snippet = d.page_content.strip().replace("\n", " ")[:400]
        print(f"[{i}] {src} | chunk:{cid} -> {snippet}...\n")

    # Build context (limited)
    context = build_context_from_docs(retrieved_docs, char_limit=CONTEXT_CHAR_LIMIT)
    # Strict prompt: answer only from context and include source citations
    prompt = (
        "You are an assistant that MUST answer only using the provided CONTEXT from documents below.\n"
        "If the answer is not present in the context, respond exactly: "
        "\"I could not find this information in the provided document.\"\n\n"
        "CONTEXT:\n"
        f"{context}\n\n"
        "QUESTION:\n"
        f"{query}\n\n"
        "Provide a concise, factual answer. Then list sources (filename and chunk number) that support the answer.\n"
        "If you include a short quote from the context, prefix it with QUOTE: and the quote.\n"
    )

    # Call LLM
    try:
        output = hf_pipe(prompt, max_new_tokens=512)
        if isinstance(output, list) and output:
            text = output[0].get("generated_text") or output[0].get("text") or str(output[0])
        else:
            text = str(output)
    except Exception as e:
        text = f"❗ LLM call failed: {e}"

    print("\n🤖 AI Answer:\n")
    # strip leading/trailing whitespace
    print(text.strip())
    print("-" * 120)
